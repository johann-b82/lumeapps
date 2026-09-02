from io import BytesIO
from types import SimpleNamespace

from docx import Document

from app.services.atr_generate_docx import (
    build_container_label,
    build_containerbeschriftung,
    grid_columns,
)


def _it(po_pos, category="Teppiche"):
    return SimpleNamespace(po_pos=po_pos, category=category)


def test_label_lines():
    delivery = SimpleNamespace(ba_auftrag="1024738", po_number="4501119979",
                               ac_programme="A350", msn="830", container_number="AK111XXX")
    items = [_it("300"), _it("050"), _it("340")]
    out = build_containerbeschriftung(delivery, items)
    doc = Document(BytesIO(out))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "BA 1024738" in text
    assert "PO 4501119979" in text
    assert "Pos. 050, 300, 340" in text  # sorted, comma-joined
    assert "A350 Teppiche MSN 830" in text
    assert "Container AK111XXX" in text


def _delivery(ba, msn, po="4501119979"):
    return SimpleNamespace(ba_auftrag=ba, po_number=po, ac_programme="A350", msn=msn,
                           container_number="AK111XXX")


def _all_paragraphs(doc):
    """Body paragraphs plus every table-cell paragraph, in document order."""
    out = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                out.extend(cell.paragraphs)
    return out


def _label_doc(n):
    return Document(BytesIO(build_container_label("C1", [
        (_delivery(f"10247{i:02d}", str(830 + i)), [_it("050")]) for i in range(n)])))


def test_container_label_lists_every_delivery():
    out = build_container_label("AK111XXX", [
        (_delivery("1024738", "830"), [_it("300"), _it("050")]),
        (_delivery("1024739", "831"), [_it("010")]),
    ])
    doc = Document(BytesIO(out))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert text.startswith("Container AK111XXX")
    assert "BA 1024738" in text and "BA 1024739" in text
    assert "Pos. 050, 300" in text and "Pos. 010" in text
    assert "A350 Teppiche MSN 830" in text and "A350 Teppiche MSN 831" in text
    assert text.count("Container ") == 1  # heading only, no per-delivery line
    assert not doc.tables  # up to two deliveries: plain single column


def test_grid_columns_grow_with_count():
    assert [grid_columns(n) for n in (3, 4, 5, 8, 9, 12, 30)] == [1, 1, 2, 2, 3, 3, 3]


def test_container_label_grid_keeps_every_block_and_font_size():
    for n, cols in ((3, 1), (7, 2), (12, 3)):
        doc = _label_doc(n)
        assert doc.paragraphs[0].text == "Container C1"
        assert len(doc.tables) == 1
        table = doc.tables[0]
        assert len(table.columns) == cols
        texts = [p.text for p in _all_paragraphs(doc)]
        for i in range(n):
            assert f"BA 10247{i:02d}" in texts
            assert f"A350 Teppiche MSN {830 + i}" in texts
        # every BA line shares one size — nothing is scaled down per count
        ba_sizes = {p.runs[0].font.size for p in _all_paragraphs(doc) if p.text.startswith("BA ")}
        assert len(ba_sizes) == 1
        assert ba_sizes.pop().pt >= 18
