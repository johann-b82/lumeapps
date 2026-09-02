"""Containerbeschriftung (.docx) generator (ATR Phase B).

Two entry points: ``build_containerbeschriftung`` (one delivery, generated
alongside the ATR) and ``build_container_label`` (all deliveries sharing one
container number, built on the fly for the deliveries list).
"""
from __future__ import annotations

from io import BytesIO
from math import ceil

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

from app.services.atr_format import normalize_po_pos

_BLACK = RGBColor(0x00, 0x00, 0x00)
_RED = RGBColor(0xC0, 0x00, 0x00)
_GREEN = RGBColor(0x00, 0x80, 0x00)
_BIG = Pt(32)      # BA + MSN line — emphasised
_NORMAL = Pt(20)

# Multi-delivery label: blocks are laid out in a borderless grid. The font
# stays fixed per column count (readable, never scaled down); more deliveries
# add columns (up to 3), beyond that the grid simply continues on page 2.
_GRID_ROWS_PER_PAGE = 4
_GRID_MAX_COLS = 3
_GRID_FONTS = {1: (Pt(24), Pt(15)), 2: (Pt(20), Pt(13)), 3: (Pt(18), Pt(12))}
_GRID_MARGIN = Cm(1.5)
_GRID_ROW_GAP = Pt(10)


def _delivery_lines(delivery, items, big=_BIG, normal=_NORMAL):
    """(text, colour, size) per line: BA black, PO/Pos red, MSN green; BA and
    the MSN line larger."""
    pos_list = sorted(
        {p for it in items if (p := (normalize_po_pos(it.po_pos) or "").strip())}
    )
    # All product groups (categories) actually shipped, in first-seen order —
    # not just carpets. Matches the ATR's section grouping.
    groups = ", ".join(dict.fromkeys(
        c for it in items if (c := (it.category or "").strip())
    ))
    return [
        (f"BA {delivery.ba_auftrag or ''}".rstrip(), _BLACK, big),
        (f"PO {delivery.po_number or ''}".rstrip(), _RED, normal),
        (f"Pos. {', '.join(pos_list)}", _RED, normal),
        (" ".join(s for s in (
            delivery.ac_programme or "", groups, f"MSN {delivery.msn or ''}".rstrip()
        ) if s), _GREEN, big),
    ]


def _landscape_doc(margin=None):
    doc = Document()
    # Landscape: switch orientation and swap page width/height.
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    if margin is not None:
        section.left_margin = section.right_margin = margin
        section.top_margin = section.bottom_margin = margin
    return doc


def _fill(paragraph, text, colour, size, compact=False):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if compact:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = size
    run.font.color.rgb = colour
    return paragraph


def _save(doc) -> bytes:
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()


def _render(lines) -> bytes:
    doc = _landscape_doc()
    for text, colour, size in lines:
        _fill(doc.add_paragraph(), text, colour, size)
    return _save(doc)


def build_containerbeschriftung(delivery, items) -> bytes:
    lines = _delivery_lines(delivery, items) + [
        (f"Container {delivery.container_number or ''}".rstrip(), _BLACK, _NORMAL),
    ]
    return _render(lines)


def grid_columns(n: int) -> int:
    """Columns for ``n`` delivery blocks: single column while it fits one page,
    then 2 or 3 columns (never more)."""
    return max(1, min(_GRID_MAX_COLS, ceil(n / _GRID_ROWS_PER_PAGE)))


def build_container_label(container_number: str, deliveries) -> bytes:
    """One label for a whole container: heading with the container number, then
    one block (BA/PO/Pos/MSN) per delivery. ``deliveries`` is a list of
    ``(delivery, items)`` pairs.

    One or two deliveries keep the large single-column look of the per-delivery
    label. From three on, the blocks go into a borderless grid with a fixed,
    readable font: one column up to four blocks, two columns up to eight, three
    columns beyond — more than twelve simply continue on the next page."""
    n = len(deliveries)
    if n <= 2:
        lines = [(f"Container {container_number}", _BLACK, _BIG)]
        for delivery, items in deliveries:
            lines.append(("", _BLACK, Pt(12)))
            lines.extend(_delivery_lines(delivery, items))
        return _render(lines)

    doc = _landscape_doc(margin=_GRID_MARGIN)
    _fill(doc.add_paragraph(), f"Container {container_number}", _BLACK, _BIG)

    cols = grid_columns(n)
    rows = ceil(n / cols)
    big, normal = _GRID_FONTS[cols]
    section = doc.sections[0]
    col_width = int((section.page_width - section.left_margin - section.right_margin) / cols)

    table = doc.add_table(rows=rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for idx, (delivery, items) in enumerate(deliveries):
        cell = table.cell(idx // cols, idx % cols)
        cell.width = col_width
        for i, (text, colour, size) in enumerate(_delivery_lines(delivery, items, big, normal)):
            p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
            _fill(p, text, colour, size, compact=True)
            if i == 0:
                p.paragraph_format.space_before = _GRID_ROW_GAP
    # Unused trailing cells keep their default width so the grid stays even.
    for cell in table.rows[-1].cells:
        cell.width = col_width
    return _save(doc)
