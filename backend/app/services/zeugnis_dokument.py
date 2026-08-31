"""Arbeitszeugnis als DOCX (editierbar) und PDF (v1.110).

Der Fließtext kommt aus ``zeugnis.abschnitte_json`` (Baukasten oder KI, danach
von HR editiert). Das DOCX wird der ACM-Endzeugnis-Vorlage nachgebaut:
Zertifizierungs-Briefkopf, Firmenbeschreibung, Aufgaben als Aufzählung, die
Beurteilungs-Absätze und zwei Unterschriften — der/die **Supervisor:in kommt
aus Personio** (abteilungsabhängig, vom Aufrufer übergeben), der/die zweite
Unterzeichner:in (HR) aus dem Ausstellerprofil. Das PDF entsteht per LibreOffice
— derselbe ``soffice``-Pfad wie Wartung/ATR/Signage.
"""
from __future__ import annotations

import asyncio
import shutil
import uuid as _uuid
from datetime import date
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from app.models.zeugnis import Zeugnis, ZeugnisAussteller

# --- ACM-Briefkopf (feste Zertifizierungen + Firmenbeschreibung) ------------
#: Nur einblenden, wenn der Aussteller die ACM ist — ein anderer Betrieb bekommt
#: seinen eigenen (schlichten) Kopf ohne diese festen Angaben.
_ACM_KENN = "aircraft cabin modification"
_ACM_ZERTIFIKATE = (
    "EASA Part 21J    EASA.21J.456",
    "EASA Part 21G    DE.21G.0170",
    "EASA Part 145    DE.145.0222",
    "DIN EN ISO 9001:2015",
    "DIN EN ISO 9100:2018",
    "DIN EN ISO 9110:2018",
)
_ACM_BESCHREIBUNG = (
    "Die Aircraft Cabin Modification GmbH ist ein nach EASA Part 21J, 21G und "
    "145 zertifizierter Betrieb. Damit qualifiziert sie sich als idealer Partner "
    "in der Luftfahrtindustrie in Sachen Entwicklung, Produktion, Instandhaltung "
    "und Überarbeitung von Flugzeuginnenausstattungen.",
    "Die Kombination von Entwicklungs-, Produktions- und Instandhaltungsbetrieb "
    "führt zu einem einzigartigen Portfolio an Produkten und Dienstleistungen, "
    "das zusätzlich ständig weiter entwickelt und ausgebaut wird.",
    "Zu unserer Produktpalette gehören die Herstellung, Überholung und Reparatur "
    "von hochwertigen Flugzeugsitzbezügen, Sicherheits- und Anschnallgurten sowie "
    "Kabinen-Equipment. Des Weiteren stellen wir Composite Verbundteile und "
    "Paneele, Crew Rest Compartments und textile Innenausstattung für Flugzeuge "
    "und Helikopter her.",
)


def _datum_kurz(d) -> str:
    return d.strftime("%d.%m.%Y") if d else ""


def _ist_acm(aussteller: ZeugnisAussteller | None) -> bool:
    return bool(aussteller and aussteller.firma and _ACM_KENN in aussteller.firma.lower())


def _titel(art: str, austritt) -> str:
    return {
        "zwischenzeugnis": "Zwischenzeugnis",
        "ausbildungszeugnis": "Ausbildungszeugnis",
        "praktikumszeugnis": "Praktikumszeugnis",
        "einfach": "Arbeitszeugnis",
    }.get(art, "Endzeugnis" if austritt else "Zeugnis")


def _tabelle_unterlinie(tabelle) -> None:
    """Dünne graue Linie unter der Kopfzeilen-Tabelle (Briefkopf-Trenner)."""
    borders = OxmlElement("w:tblBorders")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")  # 0,75 pt
    bottom.set(qn("w:space"), "0")
    bottom.set(qn("w:color"), "AAAAAA")
    borders.append(bottom)
    tabelle._tbl.tblPr.append(borders)


def _briefkopf(doc: Document, aussteller: ZeugnisAussteller | None, logo_png: bytes | None) -> None:
    """Logo (links) + Zertifizierungen (rechts) als kompakte Kopfzeile — je Seite."""
    header = doc.sections[0].header
    tabelle = header.add_table(rows=1, cols=2, width=Cm(16.0))
    tabelle.autofit = False
    links, rechts = tabelle.rows[0].cells
    links.width = Cm(7.0)
    rechts.width = Cm(9.0)
    links.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    rechts.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    lp = links.paragraphs[0]
    lp.paragraph_format.space_before = Pt(0)
    lp.paragraph_format.space_after = Pt(0)
    if logo_png:
        try:
            lp.add_run().add_picture(BytesIO(logo_png), width=Cm(4.2))
        except Exception:  # pragma: no cover - defektes Bild bricht das Zeugnis nicht
            pass

    if _ist_acm(aussteller):
        erste = True
        for zeile in _ACM_ZERTIFIKATE:
            p = rechts.paragraphs[0] if erste else rechts.add_paragraph()
            erste = False
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            pf = p.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            pf.line_spacing = 1.15
            run = p.add_run(zeile)
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        _tabelle_unterlinie(tabelle)

    # Führende Leerzeile der Kopfzeile entfernen, damit der Kopf oben sauber sitzt.
    leer = header.paragraphs[0]
    if not leer.text and leer._element.getnext() is not None:
        leer._element.getparent().remove(leer._element)


def _fliesstext(doc: Document, text: str) -> None:
    para = doc.add_paragraph(text)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.space_after = Pt(10)


def _aufgaben(doc: Document, text: str) -> None:
    """Tätigkeit: erste Zeile als Einleitung, Folgezeilen als Aufzählung."""
    zeilen = [z for z in text.splitlines() if z.strip()]
    if len(zeilen) <= 1:
        _fliesstext(doc, text.replace("\n", " ").strip())
        return
    einleit = doc.add_paragraph(zeilen[0])
    einleit.paragraph_format.space_after = Pt(4)
    for punkt in zeilen[1:]:
        doc.add_paragraph(punkt.strip(), style="List Bullet")


def _zusammen(paragraph) -> None:
    """Absatz mit dem folgenden zusammenhalten (kein Seitenumbruch dazwischen)."""
    pf = paragraph.paragraph_format
    pf.keep_with_next = True
    pf.keep_together = True


def _zeile_nicht_trennen(tabelle) -> None:
    """Tabellenzeilen nicht über einen Seitenumbruch trennen (w:cantSplit)."""
    for row in tabelle.rows:
        row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))


def _feld(paragraph, code: str) -> None:
    """Word-Feld (z. B. PAGE / NUMPAGES) einfügen — von LibreOffice beim PDF-Export ausgewertet."""
    run = paragraph.add_run()
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {code} "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def _tabelle_oberlinie(tabelle) -> None:
    """Dünne graue Linie über der Tabelle (Trenner über der Fußzeile)."""
    borders = OxmlElement("w:tblBorders")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "6")  # 0,75 pt
    top.set(qn("w:space"), "0")
    top.set(qn("w:color"), "AAAAAA")
    borders.append(top)
    tabelle._tbl.tblPr.append(borders)


def _ohne_endung(dateiname: str | None) -> str:
    name = dateiname or ""
    for ext in (".pdf", ".docx"):
        if name.lower().endswith(ext):
            return name[: -len(ext)]
    return name


def _fusszeile(doc: Document, dateiname: str | None) -> None:
    """Fußzeile je Seite: Trennlinie + Dateiname (links) + „Seite X von Y" (rechts)."""
    footer = doc.sections[0].footer
    tabelle = footer.add_table(rows=1, cols=2, width=Cm(16.0))
    tabelle.autofit = False
    _tabelle_oberlinie(tabelle)
    links_z, rechts_z = tabelle.rows[0].cells
    links_z.width = Cm(9.0)
    rechts_z.width = Cm(7.0)
    grau = RGBColor(0x88, 0x88, 0x88)

    lp = links_z.paragraphs[0]
    lp.paragraph_format.space_before = Pt(4)
    lp.paragraph_format.space_after = Pt(0)
    lr = lp.add_run(_ohne_endung(dateiname))
    lr.font.size = Pt(8)
    lr.font.color.rgb = grau

    rp = rechts_z.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp.paragraph_format.space_before = Pt(4)
    rp.paragraph_format.space_after = Pt(0)
    sr = rp.add_run("Seite ")
    sr.font.size = Pt(8)
    sr.font.color.rgb = grau
    _feld(rp, "PAGE")
    mr = rp.add_run(" von ")
    mr.font.size = Pt(8)
    mr.font.color.rgb = grau
    _feld(rp, "NUMPAGES")

    # führende Leerzeile der Fußzeile entfernen, damit sie unten bündig sitzt
    leer = footer.paragraphs[0]
    if not leer.text and leer._element.getnext() is not None:
        leer._element.getparent().remove(leer._element)


def _unterschriften(
    doc: Document,
    supervisor_name: str | None,
    supervisor_titel: str | None,
    hr_name: str | None,
    hr_titel: str | None,
) -> None:
    """Zwei Unterschriften nebeneinander: Supervisor + HR-Manager (beide aus Personio)."""
    if not supervisor_name and not hr_name:
        return
    _zusammen(doc.add_paragraph())
    _zusammen(doc.add_paragraph())  # Platz für die eigentliche Unterschrift
    tabelle = doc.add_table(rows=1, cols=2)
    _zeile_nicht_trennen(tabelle)
    for zelle, name, titel in (
        (tabelle.rows[0].cells[0], supervisor_name, supervisor_titel),
        (tabelle.rows[0].cells[1], hr_name, hr_titel),
    ):
        if not name:
            continue
        zelle.paragraphs[0].add_run(name).font.size = Pt(10)
        if titel:
            r = zelle.add_paragraph().add_run(f"– {titel} –")
            r.font.size = Pt(10)
            r.italic = True


def build_zeugnis_docx(
    zeugnis: Zeugnis,
    aussteller: ZeugnisAussteller | None,
    logo_png: bytes | None,
    *,
    supervisor_name: str | None = None,
    supervisor_titel: str | None = None,
    hr_name: str | None = None,
    hr_titel: str | None = None,
    dateiname: str | None = None,
) -> bytes:
    """Baut das Zeugnis als .docx (ACM-Endzeugnis-Vorlage) und gibt die Bytes zurück."""
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    _briefkopf(doc, aussteller, logo_png)

    # --- Titel ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(_titel(zeugnis.art, zeugnis.austritt))
    run.bold = True
    run.font.size = Pt(16)
    doc.add_paragraph()

    abschnitte = zeugnis.abschnitte_json or {}

    # --- Einleitung ---
    einleitung = str(abschnitte.get("einleitung", "") or "").strip()
    if einleitung:
        _fliesstext(doc, einleitung)

    # --- Firmenbeschreibung (nur ACM) ---
    if _ist_acm(aussteller):
        for absatz in _ACM_BESCHREIBUNG:
            _fliesstext(doc, absatz)

    # --- Tätigkeiten (als Aufzählung) ---
    taetigkeit = str(abschnitte.get("taetigkeitsbeschreibung", "") or "").strip()
    if taetigkeit:
        _aufgaben(doc, taetigkeit)
        doc.add_paragraph()

    # --- Beurteilung + Schluss ---
    for key in ("leistungsbeurteilung", "sozialverhalten", "schlussformel"):
        text = str(abschnitte.get(key, "") or "").strip()
        if text:
            _fliesstext(doc, text)

    # --- Abschluss-Block: Ort/Datum + Firma + Unterschriften (zusammenhalten) ---
    _zusammen(doc.add_paragraph())
    ort = (aussteller.standort if aussteller else None) or ""
    datum = _datum_kurz(zeugnis.ausstellungsdatum or zeugnis.austritt or date.today())
    if ort or datum:
        _zusammen(doc.add_paragraph(", ".join(x for x in (ort, datum) if x)))
    if aussteller and aussteller.firma:
        _zusammen(doc.add_paragraph())
        firma_p = doc.add_paragraph(aussteller.firma)
        firma_p.runs[0].bold = True
        _zusammen(firma_p)
    _unterschriften(doc, supervisor_name, supervisor_titel, hr_name, hr_titel)

    # --- Fußzeile: Dateiname + Seite X von Y ---
    _fusszeile(doc, dateiname)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# LibreOffice über den api-Container serialisieren (wie Wartung/ATR/Signage).
_LO_SEMAPHORE = asyncio.Semaphore(1)
_LO_TIMEOUT_S = 60


async def convert_docx_to_pdf(docx_bytes: bytes) -> bytes:
    async with _LO_SEMAPHORE:
        tempdir = Path(f"/tmp/zeugnis_{_uuid.uuid4()}")
        try:
            tempdir.mkdir(parents=True, exist_ok=True)
            src = tempdir / "zeugnis.docx"
            src.write_bytes(docx_bytes)
            profile = tempdir / "profile"
            proc = await asyncio.create_subprocess_exec(
                "soffice", "--headless", "--invisible", "--nodefault",
                "--norestore", "--nologo", "--nofirststartwizard",
                f"-env:UserInstallation=file://{profile}",
                "--convert-to", "pdf:writer_pdf_Export",
                "--outdir", str(tempdir), str(src),
                stdout=asyncio.subprocess.PIPE, stderr=asyncio.subprocess.PIPE,
            )
            try:
                _, err = await asyncio.wait_for(
                    proc.communicate(), timeout=_LO_TIMEOUT_S
                )
            except asyncio.TimeoutError as exc:
                try:
                    proc.kill()
                except ProcessLookupError:
                    pass
                raise RuntimeError("docx->pdf conversion timed out") from exc
            out = tempdir / "zeugnis.pdf"
            if proc.returncode != 0 or not out.exists():
                raise RuntimeError(
                    f"soffice pdf export failed: {err.decode('utf-8', 'replace')[-500:]}"
                )
            return out.read_bytes()
        finally:
            shutil.rmtree(tempdir, ignore_errors=True)
